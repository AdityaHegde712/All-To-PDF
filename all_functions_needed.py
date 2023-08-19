import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.section import WD_ORIENTATION
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

# create docx, set properties
def create_document(file_name):
    # Create a new document
    doc = Document()

    # Set page orientation to portrait (change to PORTRAIT)
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT

    # Get the paragraph style for the Normal style
    paragraph_style = doc.styles['Normal'].paragraph_format

    # Set page margins
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Set default font style and size
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(12)

    # Set section title font style and size
    title_style = doc.styles['Heading 1']
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(72)
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set Consolas font style for file content
    run = doc.styles['Normal'].font
    run.name = 'Consolas'

    # Disable paragraph spacing before/after
    paragraph_style.space_before = 0
    paragraph_style.space_after = 0

    # Set line spacing
    paragraph_style.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph_style.line_spacing = Inches(0.2)

    # Save the document with the specified file name
    doc.save(f"{file_name}.docx")

    return doc

# Function to add a section head to the document
def add_section_head(doc, section_name):
    doc.add_page_break()
    title_style = doc.styles['Heading 1']
    title_para = doc.add_paragraph(section_name, style=title_style)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

# Function to add non-section head subfolder information
def add_subfolder_info(doc, subfolder_list):
    if subfolder_list:
        subfolder_names = ', '.join(subfolder_list)
        doc.add_paragraph(f"{subfolder_names} also contains the following folders: {subfolder_names}.")

# Function to add file information and content
def add_file_info_and_content(doc, file_name, file_content):
    file_heading = doc.add_heading(file_name, level=1)
    file_heading.style.font.name = 'Arial'
    file_heading.style.font.size = Pt(26)

    print(file_content[0:26])
    file_content_para = doc.add_paragraph(file_content)
    file_content_para.style.font.name = 'Consolas'
    file_content_para.style.font.size = Pt(12)

    # # Add new line characters until the next page
    # while not doc.add_paragraph("").runs:
    #     pass

# Fix the folder path in process_folder
def fix_folder(bad_path, top):
    last = bad_path.rfind("\\")
    return bad_path[:last+1] + top + "\\" + bad_path[last+1:]


# Function to process a folder and its contents
def process_folder(doc, folder_path=None, top_level_folder=None):
    if folder_path is None:
        # Iterate through items in the top-level folder
        for item in os.listdir(top_level_folder):
            item_path = os.path.join(top_level_folder, item)

            if os.path.isdir(item_path):
                process_folder(doc, folder_path=item, top_level_folder=top_level_folder)
            elif os.path.isfile(item_path):
                # Ignore files at this level
                pass

        # Save the document with the top-level folder's name
        doc.save(f"{os.path.basename(top_level_folder)}.docx")
    else:
        # Convert folder path to absolute path
        folder_path = fix_folder(bad_path=os.path.abspath(folder_path), top=top_level_folder)
        print(folder_path)
        
        # Creating the section head
        section_head = folder_path[folder_path.rfind('\\')+1:]
        print(section_head)
        add_section_head(doc, section_name=section_head)

        # Iterate through items in the current folder (section head)
        for item in os.listdir(folder_path):
            item_path = os.path.join(folder_path, item)

            if os.path.isfile(item_path) and item_path.endswith('.py'):
                with open(item_path, 'rb') as file:  # Open in binary mode
                    file_content = file.read()
                # Assuming you want the content as a string, decode it with an appropriate encoding
                file_content = file_content.decode('utf-8', errors='replace')  # Use the appropriate encoding
                add_file_info_and_content(doc, item, file_content)